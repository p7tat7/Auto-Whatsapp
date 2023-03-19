from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from openpyxl import Workbook, load_workbook
import time
import docx
from tkinter import Tk
from tkinter.filedialog import askopenfilename

target = int(0)

while target != 1 and target != 2:
    try:
        target = int(input('傳送對象：\n------------------\n1.    所有會員\n2.    個別客人\n------------------\n請輸入：'))
    except: print("請輸入數字")
    if target != 1 and target != 2:
        print("請輸入有效數字")
        time.sleep(1)
if target == 1:
    wb = load_workbook('C:\\Users\\world\\World Wide Aquarium Dropbox\\WWA SHKS\\Member Data\\Member Data.xlsx')
    ws = wb.active
elif target == 2:
    filename = askopenfilename(filetypes=[("Excel", "*.xlsx"), ("All Files", "*.*")])
    wb = load_workbook(filename)
    ws = wb.active
PATH = "C:/Users/world/OneDrive/文件/chromedriver.exe"
driver = webdriver.Chrome(PATH)
driver.get('https://web.whatsapp.com')

WebDriverWait(driver, 1000).until(
       EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/div'))
    )


messageDoc = docx.Document('C:\\Users\\world\\World Wide Aquarium Dropbox\\WWA SHKS\\實用工具\\批量whatsapp message\\群發訊息.docx')
for i in range(1,ws.max_row+1):
    driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys(ws['A'+str(i)].value)
    #driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys("63393947")
    driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').send_keys(Keys.RETURN)
    time.sleep(0.5)
    #print(len(messageDoc.paragraphs))
    for j in range(len(messageDoc.paragraphs)):
        driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[1]/div/div[2]').send_keys(messageDoc.paragraphs[j].text)
        # if messageDoc.paragraphs[j].text == "":
        driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[1]/div/div[2]').send_keys(Keys.SHIFT + Keys.RETURN)
            # try:
            #     if len(messageDoc.paragraphs[j+1]) > 0:
            #         driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[1]/div/div[2]').send_keys(Keys.SHIFT+Keys.RETURN)
            # except: True
    #driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[1]/div/div[2]').send_keys('h\n\n\n\ni')
    time.sleep(0.2)
    driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[1]/div/div[2]').send_keys(Keys.RETURN)
    time.sleep(0.5)
